"""
Tests for the new extractor format.
"""
import os
import json
from docx import Document
from parser.extractor import parse_docx
from parser.schema import ParsedDocument

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "samples")
NEW_DOCX_PATH = os.path.join(SAMPLES_DIR, "sample_new_format.docx")
NEW_JSON_PATH = os.path.join(SAMPLES_DIR, "sample_new_format.json")


def create_sample_docx():
    """Creates a sample docx file for testing the new format."""
    doc = Document()
    doc.add_paragraph("# Curso: [Sample Course Name]")
    doc.add_paragraph("## Caderno: [Sample Notebook Name]")
    doc.add_paragraph("## Conteúdo Programático:")
    doc.add_paragraph("This is the programmatic content.")
    doc.add_paragraph("## Assunto 1: [Subject 1]")
    doc.add_paragraph("### Título do Slide (Teoria):")
    doc.add_paragraph("Title for Theory 1")
    doc.add_paragraph("Content for Theory 1.")
    doc.add_paragraph("### Enunciado do Exercício:")
    doc.add_paragraph("Solve the following questions.")
    doc.add_paragraph("### Questões do Exercício:")
    doc.add_paragraph("a) What is 1+1?")
    doc.add_paragraph(">2")
    doc.add_paragraph("b) What is the capital of France?")
    doc.add_paragraph(">Paris")
    doc.add_paragraph("## Questões de Concurso")
    doc.add_paragraph("### Questão 1")
    doc.add_paragraph("**Enunciado da Questão:** (CESPE/2024) This is a contest question.")
    doc.add_paragraph("**Texto:** []")
    doc.add_paragraph("### Alternativas:")
    doc.add_paragraph("- A) Option A")
    doc.add_paragraph("- B) Option B (gabarito)")
    doc.add_paragraph("- C) Option C")
    doc.save(NEW_DOCX_PATH)


def create_expected_json():
    """Creates the expected JSON output for the sample docx."""
    expected_data = {
        "courseTitle": "Sample Course Name",
        "notebookTitle": "Sample Notebook Name",
        "programmaticContent": "This is the programmatic content.",
        "subjects": [
            {
                "subjectName": "Subject 1",
                "theorySlides": [
                    {"title": "Title for Theory 1", "content": "Content for Theory 1."}
                ],
                "exercises": [
                    {
                        "statement": "Solve the following questions.",
                        "questions": [
                            {"question": "a) What is 1+1?", "answer": "2"},
                            {
                                "question": "b) What is the capital of France?",
                                "answer": "Paris",
                            },
                        ],
                    }
                ],
            }
        ],
        "contestQuestions": [
            {
                "id": 1,
                "statement": "(CESPE/2024) This is a contest question.",
                "text": "",
                "source": "CESPE/2024",
                "options": ["A) Option A", "B) Option B", "C) Option C"],
                "answer": "B",
            }
        ],
        "warnings": [],
    }
    with open(NEW_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(expected_data, f, indent=2, ensure_ascii=False)


def setup_module(module):
    """Setup test module."""
    create_sample_docx()
    create_expected_json()


def teardown_module(module):
    """Teardown test module."""
    #os.remove(NEW_DOCX_PATH)
    #os.remove(NEW_JSON_PATH)
    pass


def test_new_format_parsing():
    """
    Tests the parsing of a .docx file with the new format.
    """
    parsed_data = parse_docx(NEW_DOCX_PATH)
    
    with open(NEW_JSON_PATH, "r", encoding="utf-8") as f:
        expected_data = json.load(f)

    # Validate with Pydantic model
    validated_data = ParsedDocument(**parsed_data)
    validated_json = json.loads(validated_data.model_dump_json(by_alias=True))

    assert validated_json == expected_data
