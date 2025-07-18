"""
Creates a sample docx file for testing warnings.
"""
from docx import Document
from docx.shared import Inches

def create_warning_sample_docx():
    """
    Creates a sample docx file for testing warnings.
    """
    document = Document()

    document.add_paragraph('reta final polícia federal') # Not in all caps
    document.add_heading('CADERNO 2', level=2)

    document.add_heading('Reconhecimento de Gêneros Textuais', level=3)
    
    document.add_paragraph('Questões de Exercícios')
    
    document.add_paragraph('1) (CESPE/2024) O texto 10A2‑I é predominantemente...')
    document.add_paragraph('A) Narrativo')
    document.add_paragraph('B) Expositivo')
    document.add_paragraph('C) Descritivo')
    document.add_paragraph('D) Dissertativo')
    document.add_paragraph('E) Injuntivo')
    # Missing answer

    document.save('samples/sample2.docx')

if __name__ == '__main__':
    create_warning_sample_docx()
