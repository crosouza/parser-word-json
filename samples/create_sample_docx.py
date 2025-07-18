"""
Creates a sample docx file for testing.
"""
from docx import Document
from docx.shared import Inches

def create_sample_docx():
    """
    Creates a sample docx file for testing.
    """
    document = Document()

    document.add_heading('RETA FINAL POLÍCIA FEDERAL', level=1)
    document.add_heading('CADERNO 1', level=2)

    document.add_heading('Reconhecimento de Gêneros Textuais', level=3)
    document.add_paragraph('Interpretação')

    document.add_heading('Teoria', level=4)
    document.add_paragraph('Identificando Gêneros')
    document.add_paragraph('Estratégias de Leitura')

    document.add_paragraph('Questões de Exercícios')
    document.add_paragraph('Resolva as questões a seguir:')

    document.add_paragraph('1) (CESPE/2024) O texto 10A2‑I é predominantemente...')
    document.add_paragraph('A) Narrativo')
    document.add_paragraph('B) Expositivo')
    document.add_paragraph('C) Descritivo')
    document.add_paragraph('D) Dissertativo')
    document.add_paragraph('E) Injuntivo')
    document.add_paragraph('Gabarito: B')

    document.add_paragraph('2) Julgue o item a seguir')
    document.add_paragraph('CERTO')
    document.add_paragraph('ERRADO')
    document.add_paragraph('Gabarito: CERTO')

    document.add_heading('DIREITO ADMINISTRATIVO', level=3)
    
    document.add_heading('Teoria', level=4)
    document.add_paragraph('Atos Administrativos')

    document.add_paragraph('Questões de Exercícios')
    
    document.add_paragraph('1) (CESPE/2023) Atos administrativos são...')
    document.add_paragraph('A) Fatos')
    document.add_paragraph('B) Contratos')
    document.add_paragraph('C) Atos da administração')
    document.add_paragraph('D) Atos de direito privado')
    document.add_paragraph('E) Atos de império')
    document.add_paragraph('Gabarito: C')


    document.save('samples/sample1.docx')

if __name__ == '__main__':
    create_sample_docx()
