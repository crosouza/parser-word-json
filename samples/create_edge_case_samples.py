"""
Creates sample docx files for edge case testing.
"""
from docx import Document

def create_no_theory_doc():
    doc = Document()
    doc.add_heading("COURSE TITLE", level=1)
    doc.add_heading("NOTEBOOK 1", level=2)
    doc.add_heading("SUBJECT 1", level=3)
    doc.add_paragraph("Questões de Exercícios")
    doc.add_paragraph("1) An exercise statement.")
    doc.add_paragraph("A) Option 1")
    doc.add_paragraph("Gabarito: A")
    doc.save("samples/sample_no_theory.docx")

def create_no_exercises_doc():
    doc = Document()
    doc.add_heading("COURSE TITLE", level=1)
    doc.add_heading("NOTEBOOK 1", level=2)
    doc.add_heading("SUBJECT 1", level=3)
    doc.add_paragraph("Some theory here.")
    doc.add_heading("SUBJECT 2", level=3)
    doc.add_paragraph("More theory.")
    doc.save("samples/sample_no_exercises.docx")

def create_missing_alternatives_doc():
    doc = Document()
    doc.add_heading("COURSE TITLE", level=1)
    doc.add_heading("NOTEBOOK 1", level=2)
    doc.add_heading("SUBJECT 1", level=3)
    doc.add_paragraph("Questões de Exercícios")
    doc.add_paragraph("1) An exercise statement without options.")
    doc.add_paragraph("Gabarito: A")
    doc.save("samples/sample_missing_alternatives.docx")

def create_google_docs_export_doc():
    doc = Document()
    # Google Docs exports often lose heading styles, so we rely on ALL CAPS
    doc.add_paragraph("COURSE TITLE")
    doc.add_paragraph("NOTEBOOK 1")
    doc.add_paragraph("SUBJECT 1")
    doc.add_paragraph("Some theory.")
    doc.save("samples/sample_google_docs.docx")

def create_alternative_markers_doc():
    doc = Document()
    doc.add_heading("COURSE TITLE", level=1)
    doc.add_heading("NOTEBOOK 1", level=2)
    doc.add_heading("SUBJECT 1", level=3)
    doc.add_paragraph("Questão - Exercício")
    doc.add_paragraph("1. A statement.")
    doc.add_paragraph("(A) Option 1")
    doc.add_paragraph("[B] Option 2")
    doc.add_paragraph("C. Option 3")
    doc.add_paragraph("Alternativa Correta: B")
    doc.save("samples/sample_alt_markers.docx")

if __name__ == '__main__':
    create_no_theory_doc()
    create_no_exercises_doc()
    create_missing_alternatives_doc()
    create_google_docs_export_doc()
    create_alternative_markers_doc()
